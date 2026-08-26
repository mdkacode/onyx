"""Word rendering via python-docx.

Built for *editable* output, which drives most of the choices here. Headings
use Word's built-in named styles rather than manually-sized bold runs, bullets
use the List Bullet style, and tables use a real table style -- so Word's
navigation pane, a generated table of contents, and one-click restyling all
work in the delivered file. A visually-identical document built from manual
formatting would look the same and be miserable to edit.
"""

import io
import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from typing_extensions import override

from onyx.tools.tool_implementations.document_generation.disclaimer import (
    DISCLAIMER_TEXT,
)
from onyx.tools.tool_implementations.document_generation.models import BrandConfig
from onyx.tools.tool_implementations.document_generation.models import DocumentSpec
from onyx.tools.tool_implementations.document_generation.models import RenderedDocument
from onyx.tools.tool_implementations.document_generation.models import Section
from onyx.tools.tool_implementations.document_generation.renderers.base import (
    DocumentRenderer,
)

# Same inline markup the PDF renderer understands, so a spec renders
# equivalently in both formats.
_INLINE_TOKEN_RE = re.compile(r"(\*\*[^*\n]+?\*\*|`[^`\n]+?`)")


def _hex_to_rgb(value: str) -> RGBColor:
    """Convert a validated #RRGGBB / #RGB string to python-docx RGBColor."""
    raw = value.lstrip("#")
    if len(raw) in (3, 4):
        raw = "".join(ch * 2 for ch in raw[:3])
    return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))


def _add_inline_runs(paragraph: Paragraph, text: str) -> None:
    """Write `text` into `paragraph`, honouring **bold** and `code` markers.

    Splitting on the markup and emitting separate runs keeps the styling as
    real Word character formatting, which survives editing -- unlike leaving
    the asterisks in the text for the reader to decode.
    """
    if not text:
        return
    for token in _INLINE_TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(token)


class DocxRenderer(DocumentRenderer):
    FORMAT = "docx"
    MIME_TYPE = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    EXTENSION = "docx"

    @override
    def render(self, spec: DocumentSpec) -> RenderedDocument:
        document = Document()

        self._apply_brand(document, spec.brand)
        self._add_title_block(document, spec)
        self._add_metadata_block(document, spec)

        for section in spec.sections:
            self._add_section(document, section, spec.brand)

        self._add_disclaimer(document)
        self._set_footer(document)

        buffer = io.BytesIO()
        document.save(buffer)

        return RenderedDocument(
            content=buffer.getvalue(),
            mime_type=self.MIME_TYPE,
            extension=self.EXTENSION,
            # Word paginates at open time, so there is no page count to report
            # without rendering. Sections are the honest unit here.
            unit_count=len(spec.sections),
            unit_label="sections",
        )

    def _apply_brand(self, document: DocxDocument, brand: BrandConfig) -> None:
        """Recolor the built-in heading styles rather than styling each run.

        Editing one style restyles every heading in the document, which is the
        behaviour someone opening this in Word expects.
        """
        primary = _hex_to_rgb(brand.primary_color)
        secondary = _hex_to_rgb(brand.secondary_color)
        styles = document.styles

        for style_name, color in (
            ("Title", primary),
            ("Heading 1", primary),
            ("Heading 2", secondary),
            ("Heading 3", secondary),
        ):
            try:
                styles[style_name].font.color.rgb = color
            except KeyError:
                # A locale-specific template may not define every style; the
                # document is still valid without the recolor.
                continue

    def _add_title_block(self, document: DocxDocument, spec: DocumentSpec) -> None:
        document.add_heading(spec.title, level=0)
        if spec.subtitle:
            paragraph = document.add_paragraph(spec.subtitle)
            paragraph.style = document.styles["Subtitle"]

    def _add_metadata_block(self, document: DocxDocument, spec: DocumentSpec) -> None:
        if spec.metadata is None:
            return
        pairs = [
            ("Author", spec.metadata.author),
            ("Department", spec.metadata.department),
            ("Date", spec.metadata.date),
            ("Confidentiality", spec.metadata.confidentiality),
        ]
        present = [(label, value) for label, value in pairs if value]
        if not present:
            return
        for label, value in present:
            paragraph = document.add_paragraph()
            label_run = paragraph.add_run(f"{label}: ")
            label_run.bold = True
            paragraph.add_run(str(value))

    def _add_section(
        self, document: DocxDocument, section: Section, brand: BrandConfig
    ) -> None:
        document.add_heading(section.heading, level=1)

        if section.body:
            _add_inline_runs(document.add_paragraph(), section.body)

        for bullet in section.bullet_points:
            _add_inline_runs(
                document.add_paragraph(style="List Bullet"),
                bullet,
            )

        if section.callout:
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["Intense Quote"]
            _add_inline_runs(paragraph, section.callout)

        if section.table is not None:
            self._add_table(document, section, brand)

    def _add_table(
        self, document: DocxDocument, section: Section, brand: BrandConfig
    ) -> None:
        table_data = section.table
        if table_data is None:
            return

        table = document.add_table(rows=1, cols=len(table_data.headers))
        # A named table style keeps banding and borders editable from Word's
        # Table Design tab instead of being baked into each cell.
        try:
            table.style = document.styles["Light Grid Accent 1"]
        except KeyError:
            table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for index, header in enumerate(table_data.headers):
            cell = header_cells[index]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.color.rgb = _hex_to_rgb(brand.primary_color)

        width = len(table_data.headers)
        for row in table_data.rows:
            cells = table.add_row().cells
            # Ragged rows are common in LLM output; pad or truncate rather
            # than raising, so one malformed row can't lose the whole table.
            for index in range(width):
                cells[index].text = str(row[index]) if index < len(row) else ""

    def _add_disclaimer(self, document: DocxDocument) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(DISCLAIMER_TEXT)
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    def _set_footer(self, document: DocxDocument) -> None:
        """Repeat the notice in the page footer so it survives excerpting."""
        for doc_section in document.sections:
            footer_paragraph = doc_section.footer.paragraphs[0]
            footer_paragraph.text = DISCLAIMER_TEXT
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in footer_paragraph.runs:
                run.italic = True
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
