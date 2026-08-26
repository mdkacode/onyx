"""Tests for the document generation renderers.

The round trips here are deliberate: parsing the output back with the same
library a user's Word or Excel would use makes the assertions substantive
rather than smoke tests. The DOCX style assertions in particular encode the
"must be editable" requirement, which is easy to regress invisibly -- a
document built from manually-formatted runs looks identical and is miserable
to restyle.
"""

import csv
import io

import pytest
from docx import Document as read_docx
from openpyxl import load_workbook

from onyx.tools.tool_implementations.document_generation.disclaimer import (
    DISCLAIMER_TEXT,
)
from onyx.tools.tool_implementations.document_generation.disclaimer import (
    DISCLAIMER_TEXT_LONG,
)
from onyx.tools.tool_implementations.document_generation.models import DocumentSpec
from onyx.tools.tool_implementations.document_generation.models import Section
from onyx.tools.tool_implementations.document_generation.models import TableData
from onyx.tools.tool_implementations.document_generation.renderers import (
    available_formats,
)
from onyx.tools.tool_implementations.document_generation.renderers import get_renderer
from onyx.tools.tool_implementations.document_generation.renderers.docx import (
    DocxRenderer,
)
from onyx.tools.tool_implementations.document_generation.renderers.tabular import (
    CsvRenderer,
)
from onyx.tools.tool_implementations.document_generation.renderers.tabular import (
    sanitize_cell,
)
from onyx.tools.tool_implementations.document_generation.renderers.tabular import (
    XlsxRenderer,
)


def _spec(**overrides: object) -> DocumentSpec:
    defaults: dict[str, object] = {
        "title": "Q3 Fleet Report",
        "subtitle": "Prepared for the operations team",
        "sections": [
            Section(
                heading="Executive Summary",
                body="Utilisation rose **12%** against the `Q2` baseline.",
                bullet_points=["Uptime improved", "Charging costs fell"],
                callout="Two depots remain over capacity.",
            ),
            Section(
                heading="Fleet Metrics",
                table=TableData(
                    headers=["Depot", "Buses", "Uptime"],
                    rows=[["Pune", "40", "97%"], ["Nashik", "25", "94%"]],
                ),
            ),
        ],
    }
    defaults.update(overrides)
    return DocumentSpec(**defaults)  # ty: ignore[invalid-argument-type]


class TestDocxIsEditable:
    """The DOCX must use real Word styles, not hand-rolled formatting."""

    def test_headings_use_named_word_styles(self) -> None:
        rendered = DocxRenderer().render(_spec())
        document = read_docx(io.BytesIO(rendered.content))

        styles_used = {
            p.style.name
            for p in document.paragraphs
            if p.text.strip() and p.style is not None
        }

        # Named heading styles are what drive Word's navigation pane and any
        # generated table of contents. Bold runs would look the same and do
        # neither.
        assert "Title" in styles_used
        assert "Heading 1" in styles_used

    def test_bullets_use_list_style(self) -> None:
        rendered = DocxRenderer().render(_spec())
        document = read_docx(io.BytesIO(rendered.content))

        bullet_paragraphs = [
            p
            for p in document.paragraphs
            if p.style is not None and p.style.name == "List Bullet"
        ]
        assert len(bullet_paragraphs) == 2
        assert bullet_paragraphs[0].text == "Uptime improved"

    def test_inline_markup_becomes_real_formatting(self) -> None:
        """**bold** must become a bold run, not literal asterisks."""
        rendered = DocxRenderer().render(_spec())
        document = read_docx(io.BytesIO(rendered.content))

        body = next(p for p in document.paragraphs if "Utilisation rose" in p.text)
        assert "**" not in body.text
        assert "`" not in body.text
        assert any(run.bold and run.text == "12%" for run in body.runs)

    def test_table_content_round_trips(self) -> None:
        rendered = DocxRenderer().render(_spec())
        document = read_docx(io.BytesIO(rendered.content))

        assert len(document.tables) == 1
        table = document.tables[0]
        assert [c.text for c in table.rows[0].cells] == ["Depot", "Buses", "Uptime"]
        assert [c.text for c in table.rows[1].cells] == ["Pune", "40", "97%"]

    def test_ragged_rows_are_padded_not_fatal(self) -> None:
        """LLM output routinely has short rows; losing the table would be worse."""
        spec = _spec(
            sections=[
                Section(
                    heading="Partial",
                    table=TableData(
                        headers=["A", "B", "C"],
                        rows=[["1"], ["1", "2", "3", "4"]],
                    ),
                )
            ]
        )
        rendered = DocxRenderer().render(spec)
        document = read_docx(io.BytesIO(rendered.content))

        table = document.tables[0]
        assert [c.text for c in table.rows[1].cells] == ["1", "", ""]
        assert [c.text for c in table.rows[2].cells] == ["1", "2", "3"]


class TestFormulaInjection:
    """LLM-authored cells are executed by Excel unless neutralized."""

    @pytest.mark.parametrize(
        "dangerous",
        [
            '=HYPERLINK("http://evil","click")',
            "+1+1",
            "@SUM(A1:A9)",
            "=cmd|'/c calc'!A0",
        ],
    )
    def test_formula_triggers_are_neutralized(self, dangerous: str) -> None:
        assert sanitize_cell(dangerous).startswith("'")

    @pytest.mark.parametrize("safe", ["-5", "-3.14", "42", "Pune", ""])
    def test_ordinary_data_is_untouched(self, safe: str) -> None:
        """Negative numbers start with a trigger char but are data, not formulas."""
        assert sanitize_cell(safe) == safe

    def test_csv_output_is_sanitized_end_to_end(self) -> None:
        spec = _spec(
            sections=[
                Section(
                    heading="Payload",
                    table=TableData(headers=["col"], rows=[["=1+1"]]),
                )
            ]
        )
        rendered = CsvRenderer().render(spec)
        rows = list(csv.reader(io.StringIO(rendered.content.decode("utf-8-sig"))))
        assert rows[1] == ["'=1+1"]

    def test_xlsx_output_is_sanitized_end_to_end(self) -> None:
        spec = _spec(
            sections=[
                Section(
                    heading="Payload",
                    table=TableData(headers=["col"], rows=[["=1+1"]]),
                )
            ]
        )
        rendered = XlsxRenderer().render(spec)
        sheet = load_workbook(io.BytesIO(rendered.content)).active
        assert sheet is not None
        assert sheet["A2"].value == "'=1+1"  # ty: ignore[not-subscriptable]


class TestTabularBehaviour:
    def test_csv_exports_first_table_and_reports_the_rest(self) -> None:
        rendered = CsvRenderer().render(_spec())
        rows = list(csv.reader(io.StringIO(rendered.content.decode("utf-8-sig"))))

        assert rows[0] == ["Depot", "Buses", "Uptime"]
        assert rows[1] == ["Pune", "40", "97%"]
        # Only one table in the fixture, so nothing was dropped.
        assert rendered.notes == []

    def test_csv_warns_when_tables_are_dropped(self) -> None:
        spec = _spec(
            sections=[
                Section(heading="One", table=TableData(headers=["a"], rows=[["1"]])),
                Section(heading="Two", table=TableData(headers=["b"], rows=[["2"]])),
            ]
        )
        rendered = CsvRenderer().render(spec)

        assert rendered.notes, "dropping a table silently would lose user data"
        assert "first of 2" in rendered.notes[0]
        assert "xlsx" in rendered.notes[0]

    def test_xlsx_gives_each_table_its_own_sheet(self) -> None:
        spec = _spec(
            sections=[
                Section(heading="Depots", table=TableData(headers=["a"], rows=[["1"]])),
                Section(heading="Routes", table=TableData(headers=["b"], rows=[["2"]])),
            ]
        )
        rendered = XlsxRenderer().render(spec)
        workbook = load_workbook(io.BytesIO(rendered.content))

        assert workbook.sheetnames == ["Depots", "Routes"]
        assert rendered.unit_count == 2

    def test_xlsx_sheet_names_are_made_legal_and_unique(self) -> None:
        """Excel rejects several characters and caps names at 31 chars."""
        spec = _spec(
            sections=[
                Section(heading="A/B:C", table=TableData(headers=["x"], rows=[["1"]])),
                Section(heading="A/B:C", table=TableData(headers=["x"], rows=[["2"]])),
            ]
        )
        rendered = XlsxRenderer().render(spec)
        names = load_workbook(io.BytesIO(rendered.content)).sheetnames

        assert len(set(names)) == 2
        assert not any(c in name for name in names for c in r"[]:*?/\\")

    @pytest.mark.parametrize("renderer", [CsvRenderer(), XlsxRenderer()])
    def test_tabular_format_without_a_table_fails_clearly(
        self, renderer: object
    ) -> None:
        """The message is written for the LLM to recover from, not just to log."""
        spec = _spec(sections=[Section(heading="Prose only", body="No table here.")])

        with pytest.raises(ValueError) as excinfo:
            renderer.render(spec)  # ty: ignore[unresolved-attribute]

        message = str(excinfo.value)
        assert "table" in message.lower()
        assert "docx" in message or "pdf" in message


class TestDisclaimer:
    """Provenance must survive the file leaving Onyx."""

    def test_docx_carries_the_disclaimer_in_body_and_footer(self) -> None:
        rendered = DocxRenderer().render(_spec())
        document = read_docx(io.BytesIO(rendered.content))

        assert any(DISCLAIMER_TEXT in p.text for p in document.paragraphs)
        footers = [p.text for s in document.sections for p in s.footer.paragraphs]
        assert any(DISCLAIMER_TEXT in text for text in footers)

    def test_csv_carries_the_disclaimer(self) -> None:
        rendered = CsvRenderer().render(_spec())
        assert DISCLAIMER_TEXT_LONG in rendered.content.decode("utf-8-sig")

    def test_xlsx_carries_the_disclaimer(self) -> None:
        rendered = XlsxRenderer().render(_spec())
        sheet = load_workbook(io.BytesIO(rendered.content)).active
        assert sheet is not None
        values = [
            c.value
            for row in sheet.iter_rows()  # ty: ignore[unresolved-attribute]
            for c in row
        ]
        assert DISCLAIMER_TEXT_LONG in values


class TestRegistry:
    def test_unavailable_format_is_rejected_with_alternatives(self) -> None:
        with pytest.raises(ValueError) as excinfo:
            get_renderer("pages")
        assert "Available formats" in str(excinfo.value)

    def test_pure_python_formats_are_always_available(self) -> None:
        """PDF may be absent (WeasyPrint system libs); these never are."""
        formats = available_formats()
        assert {"docx", "csv", "xlsx"}.issubset(set(formats))

    @pytest.mark.skipif(
        "pdf" not in available_formats(),
        reason="WeasyPrint system libraries not installed in this environment",
    )
    def test_pdf_renders_when_weasyprint_is_present(self) -> None:
        rendered = get_renderer("pdf").render(_spec())
        assert rendered.content.startswith(b"%PDF")
        assert rendered.unit_count >= 1
